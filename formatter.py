import os
import io
import re
import qrcode
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx2pdf import convert

# Importa configurações centraliz adas
from config import (
    OUTPUT_DIR, DOCUMENT_WIDTH_CM, DOCUMENT_HEIGHT_CM, MARGIN_CM,
    DEFAULT_FONT, DEFAULT_FONT_SIZE, HEADING_FONT, HEADING_FONT_SIZE,
    BOX_FONT_SIZE, COLOR_NORMAL, COLOR_ALERT, COLOR_ERROR, COLOR_BOX_BG,
    FORMATTING_TAGS
)
from logger import logger
from exceptions import FormattingException, PDFConversionError


def set_p_border_and_shading(p, fill_color: str = COLOR_BOX_BG) -> None:
    """
    Define bordas superior e inferior, e cor de fundo de um parágrafo via OXML.
    Usado para simular um box profissional de livro médico.
    
    Args:
        p: Parágrafo do documento
        fill_color: Cor de fundo em hexadecimal (default: F2F2F2)
    """
    pPr = p._p.get_or_add_pPr()
    
    # Shading (Fundo)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)
    
    # Borders (Bordas)
    pbdr = OxmlElement('w:pBdr')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12') # 1.5 pt
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), '000000')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')
    
    pbdr.append(top)
    pbdr.append(bottom)
    
    pPr.append(pbdr)


def add_page_number(run) -> None:
    """
    Adiciona o campo de número de página numérico usando OXML.
    
    Args:
        run: Run do documento onde adicionar número de página
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def sanitize_text_line(line: str) -> str:
    """Remove marcas de formatação indesejadas da linha final."""
    sanitized = line.replace('**', '')
    sanitized = re.sub(r'^[\s]*[#*]+\s*', '', sanitized)
    sanitized = sanitized.replace('*', '')
    sanitized = sanitized.replace('#', '')
    return sanitized.strip()


def _to_superscript(number: int) -> str:
    superscript_map = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹'
    }
    return ''.join(superscript_map[d] for d in str(number))


def renumber_citations_and_format_superscript(text: str) -> str:
    """Renomeia referências por ordem de ocorrência e converte para sobrescrito."""
    citation_map = {}
    counter = {'next': 1}

    def replace(match):
        original = match.group(1)
        if original not in citation_map:
            citation_map[original] = counter['next']
            counter['next'] += 1
        return _to_superscript(citation_map[original])

    return re.sub(r'\[(\d+)\]', replace, text)


def generate_formatted_docx(ai_text: str, chapter_name: str) -> str:
    """
    Lê o texto gerado pela IA linha por linha, aplica estilos e tags, 
    gera QR Codes para os links, e salva o documento Word final formatado.
    
    Args:
        ai_text: Texto processado pela IA com tags de formatação
        chapter_name: Nome do capítulo (para nome do arquivo)
        
    Returns:
        Caminho completo do arquivo DOCX gerado
        
    Raises:
        FormattingException: Se falhar ao gerar ou salvar documento
    """
    ai_text = renumber_citations_and_format_superscript(ai_text)
    try:
        logger.info(f"Gerando DOCX formatado para: {chapter_name}")
        
        doc = Document()
        
        # Configura primeira seção (tamanho de livro: 17 x 24 cm, margens 1.5 cm)
        section = doc.sections[0]
        section.page_width = Cm(DOCUMENT_WIDTH_CM)
        section.page_height = Cm(DOCUMENT_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        
        logger.debug("Configurações de página definidas")
        
        # Configura layout de Duas Colunas manipulando o XML da seção
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), '2')
            cols[0].set(qn('w:space'), str(int(Cm(0.5).twips)))
        else:
            new_cols = OxmlElement('w:cols')
            new_cols.set(qn('w:num'), '2')
            new_cols.set(qn('w:space'), str(int(Cm(0.5).twips)))
            sectPr.append(new_cols)
            
        # Insere numeração de páginas no rodapé
        footer = section.footer
        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = p_footer.add_run()
        add_page_number(run_footer)
        
        # Configura o estilo Normal (Fonte Cambria/Times New Roman, Tamanho 10) para o documento padrão
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing = 1.0
        
        # Configura Títulos (Heading 1) para Arial 14pt Preto e Negrito
        style_h1 = doc.styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.name = 'Arial'
        font_h1.size = Pt(14)
        font_h1.color.rgb = RGBColor(0, 0, 0)
        font_h1.bold = True
        style_h1.paragraph_format.space_before = Pt(12)
        style_h1.paragraph_format.space_after = Pt(4)
        
        lines = ai_text.split('\n')
        
        current_mode = "DEFAULT"
        current_box_p = None
        resumo_lines = 0
        
        for line in lines:
            stripped_line = line.strip()
            
            # Tratamento para linhas vazias
            if not stripped_line:
                if current_mode == "DEFAULT":
                    doc.add_paragraph("")
                elif current_mode == "BOX_RESUMO" and current_box_p is not None:
                    # Preserve o box mas não adicione parágrafos extras vazios
                    pass
                # Sai de blocos de formatação única (exceto resumo e links) se houver quebra de parágrafo
                if current_mode in ["BOX_RECOMENDACAO", "BOX_ATENCAO", "SUGESTAO_EDICAO"]:
                    current_mode = "DEFAULT"
                continue
                
            # Identificação de Tags para mudança de estado e aplicação de estilo
            if "[BOX_RESUMO]" in stripped_line:
                current_mode = "BOX_RESUMO"
                resumo_lines = 0
                current_box_p = doc.add_paragraph()
                set_p_border_and_shading(current_box_p, COLOR_BOX_BG)
                
                run_title = current_box_p.add_run("PONTOS IMPORTANTES\n")
                run_title.bold = True
                run_title.font.name = 'Arial'
                run_title.font.size = Pt(BOX_FONT_SIZE)
                
                clean_text = sanitize_text_line(stripped_line.replace("[BOX_RESUMO]", "").strip())
                if clean_text:
                    if not clean_text.startswith("-"):
                        clean_text = f"- {clean_text}"
                    if resumo_lines < 5:
                        run_text = current_box_p.add_run(clean_text)
                        run_text.font.name = 'Arial'
                        run_text.font.size = Pt(BOX_FONT_SIZE)
                        resumo_lines += 1
                continue
                
            elif "[BOX_RECOMENDACAO]" in stripped_line:
                current_mode = "BOX_RECOMENDACAO"
                clean_text = sanitize_text_line(stripped_line.replace("[BOX_RECOMENDACAO]", "").strip())
                if clean_text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.right_indent = Cm(0.5)
                    run = p.add_run(clean_text)
                    run.bold = True
                    run.italic = True
                continue
                
            elif "[BOX_ATENCAO]" in stripped_line:
                current_mode = "BOX_ATENCAO"
                clean_text = sanitize_text_line(stripped_line.replace("[BOX_ATENCAO]", "").strip())
                if clean_text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.right_indent = Cm(0.5)
                    run = p.add_run(clean_text)
                    run.bold = True
                    run.italic = True
                    run.font.color.rgb = RGBColor(*COLOR_ALERT)
                continue
                
            elif "[SUGESTAO_EDICAO]" in stripped_line:
                current_mode = "SUGESTAO_EDICAO"
                clean_text = sanitize_text_line(stripped_line.replace("[SUGESTAO_EDICAO]", "").strip())
                if clean_text:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[SUGESTÃO DE EDIÇÃO]: {clean_text}")
                    run.font.color.rgb = RGBColor(*COLOR_ERROR)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                continue
                
            elif "[LINKS_ATUALIZACAO]" in stripped_line:
                current_mode = "LINKS_ATUALIZACAO"
                p = doc.add_paragraph()
                run = p.add_run("Links para Atualização:")
                run.bold = True
                
                clean_text = sanitize_text_line(stripped_line.replace("[LINKS_ATUALIZACAO]", "").strip())
                if not clean_text:
                    continue
                stripped_line = clean_text
                
            # Processamento contínuo das linhas baseado no estado atual
            if current_mode == "BOX_RESUMO":
                if current_box_p:
                    clean_text = sanitize_text_line(stripped_line)
                    if not clean_text:
                        continue
                    if not clean_text.startswith("-"):
                        clean_text = f"- {clean_text}"
                    if resumo_lines < 5:
                        run_text = current_box_p.add_run("\n" + clean_text)
                        run_text.font.name = 'Arial'
                        run_text.font.size = Pt(BOX_FONT_SIZE)
                        resumo_lines += 1
                    continue
            
            elif current_mode in ["BOX_RECOMENDACAO", "BOX_ATENCAO"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)
                run = p.add_run(sanitize_text_line(stripped_line))
                run.bold = True
                run.italic = True
                if current_mode == "BOX_ATENCAO":
                    run.font.color.rgb = RGBColor(*COLOR_ALERT)
                continue
            
            elif current_mode == "SUGESTAO_EDICAO":
                p = doc.add_paragraph()
                run = p.add_run(sanitize_text_line(stripped_line))
                run.font.color.rgb = RGBColor(*COLOR_ERROR)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                continue
                
            elif current_mode == "LINKS_ATUALIZACAO":
                # Extração de URLs via Regex
                urls = re.findall(r'(https?://[^\s]+)', stripped_line)
                
                if urls:
                    for url in urls:
                        # Gera o QR Code em memória usando BytesIO
                        qr = qrcode.QRCode(box_size=4, border=4)
                        qr.add_data(url)
                        qr.make(fit=True)
                        img = qr.make_image(fill_color="black", back_color="white")
                        
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        # Insere a imagem centralizada no documento
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        run_img.add_picture(img_byte_arr, width=Inches(1.2))
                        
                        # Insere a legenda curta (a própria URL em itálico ou descrição caso haja texto)
                        texto_descritivo = stripped_line.replace(url, "").strip(" -*")
                        texto_descritivo = sanitize_text_line(texto_descritivo)
                        legenda = f"{texto_descritivo}\n{url}" if texto_descritivo else url
                        
                        p_legenda = doc.add_paragraph(legenda)
                        p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_legenda.runs[0].font.size = Pt(9)
                        p_legenda.runs[0].font.italic = True
                else:
                    # Linhas descritivas de links sem URLs identificadas
                    doc.add_paragraph(sanitize_text_line(stripped_line))
                continue

            else:
                # Modo DEFAULT (Texto padrão e Títulos)
                clean_for_check = sanitize_text_line(stripped_line)
                
                # Heurística para detecção de títulos
                is_heading = (
                    len(clean_for_check) > 0 
                    and len(clean_for_check) < 60 
                    and not clean_for_check.endswith('.') 
                    and not clean_for_check.endswith(':')
                    and not clean_for_check.startswith('-')
                    and not clean_for_check.startswith('*')
                    and not clean_for_check.startswith('[')
                )
                
                if is_heading:
                    # Limpa marcadores Markdown (#) caso a IA gere, e define como título
                    heading_text = clean_for_check.lstrip('#').strip()
                    doc.add_heading(heading_text, level=1)
                else:
                    doc.add_paragraph(clean_for_check)
                continue
                

        # Sanitização do nome do arquivo (remove extensão se presente e caracteres inválidos)
        base_name = os.path.splitext(chapter_name)[0] if '.' in chapter_name else chapter_name
        safe_chapter_name = re.sub(r'[\\/*?:"<>|]', "", base_name)
        safe_chapter_name = safe_chapter_name.replace(" ", "_")
        
        output_filename = f"Capitulo_{safe_chapter_name}_Revisado.docx"
        
        doc.save(output_filename)
        logger.info(f"Documento DOCX salvo: {output_filename}")
        return output_filename
        
    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        raise FormattingException(f"Falha ao gerar documento: {e}")


def convert_to_pdf(docx_path: str) -> str:
    """
    Converte o arquivo DOCX especificado em PDF usando o docx2pdf.
    
    Args:
        docx_path: Caminho do arquivo DOCX
        
    Returns:
        Caminho do arquivo PDF gerado
        
    Raises:
        PDFConversionError: Se falhar na conversão
    """
    try:
        logger.info(f"Convertendo para PDF: {docx_path}")
        pdf_path = docx_path.replace(".docx", ".pdf")
        convert(docx_path, pdf_path)
        logger.info(f"PDF gerado com sucesso: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.error(f"Erro ao converter para PDF: {e}")
        raise PDFConversionError(f"Falha ao converter para PDF: {e}")
